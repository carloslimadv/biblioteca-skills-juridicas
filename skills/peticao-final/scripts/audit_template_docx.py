#!/usr/bin/env python3
"""Audit the peticao-final DOCX template formatting."""

from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

from defusedxml import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_INPUT = ROOT / "assets" / "modelo-peticao.docx"
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
WP_NS = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
FOOTER_LINES = (
    "Rua Esmeraldas, n. 118, Bairro Brasília, Feira de Santana/BA",
    "henriquesantana@yahoo.com.br",
    "(75) 99249-1844",
)


def close(value, expected, tolerance=1000) -> bool:
    return abs(int(value) - int(expected)) <= tolerance


def report(checks: list[tuple[str, bool]]) -> int:
    failed = [name for name, ok in checks if not ok]
    for name, ok in checks:
        print(f"{'OK' if ok else 'FAIL'} {name}")
    if failed:
        print("\nTemplate audit failed:")
        for name in failed:
            print(f"- {name}")
        return 1
    print("\nTemplate audit passed.")
    return 0


def xml_checks(path: Path) -> list[tuple[str, bool]]:
    checks: list[tuple[str, bool]] = []
    with ZipFile(path) as package:
        names = package.namelist()
        settings = ET.fromstring(package.read("word/settings.xml"))
        checks.append(("hyphenation enabled", settings.find(f".//{W_NS}autoHyphenation") is not None))

        header_names = [name for name in names if re.fullmatch(r"word/header\d+\.xml", name)]
        header_has_logo = False
        header_right_aligned = False
        for name in header_names:
            root = ET.fromstring(package.read(name))
            header_has_logo = header_has_logo or root.find(f".//{W_NS}drawing") is not None
            for paragraph in root.iter(f"{W_NS}p"):
                jc = paragraph.find(f"./{W_NS}pPr/{W_NS}jc")
                header_right_aligned = header_right_aligned or (jc is not None and jc.attrib.get(f"{W_NS}val") == "right")
        checks.append(("header logo present", header_has_logo))
        checks.append(("header logo right aligned", header_right_aligned))

        footer_names = [name for name in names if re.fullmatch(r"word/footer\d+\.xml", name)]
        footer_text = ""
        footer_sizes: list[str] = []
        footer_colors: list[str] = []
        footer_right_aligned = False
        for name in footer_names:
            root = ET.fromstring(package.read(name))
            footer_text += "\n".join(text.text or "" for text in root.iter(f"{W_NS}t"))
            for paragraph in root.iter(f"{W_NS}p"):
                jc = paragraph.find(f"./{W_NS}pPr/{W_NS}jc")
                footer_right_aligned = footer_right_aligned or (jc is not None and jc.attrib.get(f"{W_NS}val") == "right")
            for run in root.iter(f"{W_NS}r"):
                text = "".join(node.text or "" for node in run.iter(f"{W_NS}t"))
                if not text.strip():
                    continue
                size = run.find(f"./{W_NS}rPr/{W_NS}sz")
                color = run.find(f"./{W_NS}rPr/{W_NS}color")
                if size is not None:
                    footer_sizes.append(size.attrib.get(f"{W_NS}val", ""))
                if color is not None:
                    footer_colors.append(color.attrib.get(f"{W_NS}val", ""))
        checks.append(("footer right aligned", footer_right_aligned))
        for line in FOOTER_LINES:
            checks.append((f"footer text: {line}", line in footer_text))
        checks.append(("footer font 7 pt", bool(footer_sizes) and all(size == "14" for size in footer_sizes)))
        checks.append(("footer color gray 7F7F7F", bool(footer_colors) and all(color.upper() == "7F7F7F" for color in footer_colors)))
    return checks


def audit(path: Path = DEFAULT_INPUT) -> int:
    doc = Document(path)
    section = doc.sections[0]
    normal = doc.styles["Normal"]
    normal_pf = normal.paragraph_format

    checks: list[tuple[str, bool]] = [
        ("file exists", path.exists()),
        ("page width A4", close(section.page_width, Cm(21))),
        ("page height A4", close(section.page_height, Cm(29.7))),
        ("top margin 2.5 cm", close(section.top_margin, Cm(2.5))),
        ("bottom margin 2.5 cm", close(section.bottom_margin, Cm(2.5))),
        ("left margin 3.5 cm", close(section.left_margin, Cm(3.5))),
        ("right margin 3.5 cm", close(section.right_margin, Cm(3.5))),
        ("header 0 cm", close(section.header_distance, Cm(0))),
        ("footer 0 cm", close(section.footer_distance, Cm(0))),
        ("Normal font Helvetica", normal.font.name == "Helvetica"),
        ("Normal font 12 pt", close(normal.font.size, Pt(12))),
        ("Normal justified", normal_pf.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY),
        ("Normal line spacing 1.2", normal_pf.line_spacing == 1.2),
        ("Normal space before 0", close(normal_pf.space_before, Pt(0))),
        ("Normal space after 8", close(normal_pf.space_after, Pt(8))),
        ("Normal left indent 0", close(normal_pf.left_indent, Cm(0))),
        ("Normal right indent 0", close(normal_pf.right_indent, Cm(0))),
        ("Normal first line indent 0", close(normal_pf.first_line_indent, Cm(0))),
    ]

    required_styles = {
        "Peticao Titulo Principal": {"size": 12, "bold": True, "italic": False, "left": 0.0, "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY},
        "Peticao Subtitulo": {"size": 12, "bold": False, "italic": True, "left": 0.0, "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY},
        "Peticao Citacao": {"size": 11, "bold": False, "italic": False, "left": 2.5, "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY},
        "Peticao Pedido": {"size": 12, "bold": False, "italic": False, "left": 0.0, "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY},
        "Peticao Assinatura Nome": {"size": 12, "bold": True, "italic": False, "left": 0.0, "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY},
        "Peticao Assinatura OAB": {"size": 11, "bold": False, "italic": False, "left": 0.0, "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY},
    }
    for style_name, expected in required_styles.items():
        exists = style_name in doc.styles
        checks.append((f"style exists: {style_name}", exists))
        if not exists:
            continue
        style = doc.styles[style_name]
        checks.append((f"{style_name} font Helvetica", style.font.name == "Helvetica"))
        checks.append((f"{style_name} size {expected['size']} pt", close(style.font.size, Pt(expected["size"]))))
        checks.append((f"{style_name} bold", bool(style.font.bold) == expected["bold"]))
        checks.append((f"{style_name} italic", bool(style.font.italic) == expected["italic"]))
        checks.append((f"{style_name} left indent", close(style.paragraph_format.left_indent, Cm(expected["left"]))))
        checks.append((f"{style_name} justified", style.paragraph_format.alignment == expected["alignment"]))

    all_text = "\n".join(p.text for p in doc.paragraphs)
    request_lines = [p.text.strip() for p in doc.paragraphs if re.match(r"^[a-z]\)", p.text.strip())]
    checks.extend(
        [
            ("signature name present", "Carlos Henrique Santana Lima" in all_text),
            ("signature OAB present", "OAB 60.427/BA" in all_text),
            ("manual request a)", any(line.startswith("a)") for line in request_lines)),
            ("manual request b)", any(line.startswith("b)") for line in request_lines)),
            ("manual request c)", any(line.startswith("c)") for line in request_lines)),
        ]
    )
    checks.extend(xml_checks(path))
    return report(checks)


def main() -> int:
    import argparse

    parser = argparse.ArgumentParser(description="Audit the peticao-final DOCX template.")
    parser.add_argument("path", nargs="?", type=Path, default=DEFAULT_INPUT)
    args = parser.parse_args()
    return audit(args.path)


if __name__ == "__main__":
    raise SystemExit(main())
