#!/usr/bin/env python3
"""Generate the deterministic DOCX template for the peticao-final skill."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = ROOT / "assets" / "modelo-peticao.docx"
LOGO = ROOT / "assets" / "logo-carlos-lima.png"
FOOTER_LINES = (
    "Rua Esmeraldas, n. 118, Bairro Brasília, Feira de Santana/BA",
    "henriquesantana@yahoo.com.br",
    "(75) 99249-1844",
)


def set_run_font(run, *, size: int = 12, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Helvetica"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(key), "Helvetica")


def set_style_font(style, *, size: int = 12, bold: bool = False, italic: bool = False) -> None:
    style.font.name = "Helvetica"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(key), "Helvetica")


def configure_paragraph_format(style, *, left_cm: float = 0.0, size_after: int = 8) -> None:
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.2
    pf.space_before = Pt(0)
    pf.space_after = Pt(size_after)
    pf.left_indent = Cm(left_cm)
    pf.right_indent = Cm(0)
    pf.first_line_indent = Cm(0)


def set_language_pt_br(doc: Document) -> None:
    settings = doc.settings.element
    lang = settings.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        settings.append(lang)
    lang.set(qn("w:val"), "pt-BR")
    lang.set(qn("w:eastAsia"), "pt-BR")
    lang.set(qn("w:bidi"), "pt-BR")


def enable_hyphenation(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:autoHyphenation")) is None:
        settings.append(OxmlElement("w:autoHyphenation"))


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(3.5)
    section.header_distance = Cm(0)
    section.footer_distance = Cm(0)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    normal = doc.styles["Normal"]
    set_style_font(normal, size=12)
    configure_paragraph_format(normal)
    set_language_pt_br(doc)
    enable_hyphenation(doc)


def configure_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    footer = section.footer

    for container in (header, footer):
        for paragraph in container.paragraphs:
            paragraph.clear()

    header.paragraphs[0].paragraph_format.space_after = Pt(0)
    header.add_paragraph()
    logo_paragraph = header.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo_paragraph.paragraph_format.space_after = Pt(0)
    if LOGO.exists():
        logo_paragraph.add_run().add_picture(str(LOGO), width=Cm(2.51))

    footer.paragraphs[0].paragraph_format.line_spacing = 1.2
    footer.paragraphs[0].paragraph_format.space_after = Pt(0)
    for line in FOOTER_LINES:
        paragraph = footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.line_spacing = 1.2
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        set_run_font(run, size=7)
        run.font.color.rgb = RGBColor(127, 127, 127)
    footer.add_paragraph()


def ensure_style(doc: Document, name: str, *, size: int, bold: bool = False, italic: bool = False, left_cm: float = 0.0) -> None:
    styles = doc.styles
    if name in styles:
        style = styles[name]
    else:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles["Normal"]
    set_style_font(style, size=size, bold=bold, italic=italic)
    configure_paragraph_format(style, left_cm=left_cm)


def add_paragraph(doc: Document, text: str = "", style: str = "Normal"):
    paragraph = doc.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_manual_request(doc: Document, marker: str, text_before: str, keyword: str, text_after: str) -> None:
    p = doc.add_paragraph(style="Peticao Pedido")
    marker_run = p.add_run(marker)
    set_run_font(marker_run, bold=True)
    p.add_run(" ")
    before = p.add_run(text_before)
    set_run_font(before)
    key = p.add_run(keyword)
    set_run_font(key, italic=True)
    after = p.add_run(text_after)
    set_run_font(after)


def build_template(output: Path = DEFAULT_OUTPUT) -> Path:
    doc = Document()
    configure_document(doc)
    configure_header_footer(doc)

    ensure_style(doc, "Peticao Titulo Principal", size=12, bold=True)
    ensure_style(doc, "Peticao Subtitulo", size=12, italic=True)
    ensure_style(doc, "Peticao Citacao", size=11, left_cm=2.5)
    ensure_style(doc, "Peticao Pedido", size=12)
    ensure_style(doc, "Peticao Assinatura Nome", size=12, bold=True)
    ensure_style(doc, "Peticao Assinatura OAB", size=11)

    add_paragraph(doc, "AO JUIZO COMPETENTE", "Peticao Titulo Principal")
    add_paragraph(doc)
    add_paragraph(doc)
    intro = doc.add_paragraph(style="Normal")
    r = intro.add_run("[NOME DA PARTE]")
    set_run_font(r, bold=True)
    r = intro.add_run(", ja qualificada ou a qualificar, vem, por seu advogado, apresentar a presente ")
    set_run_font(r)
    r = intro.add_run("[AÇÃO DE CONHECIMENTO PELO RITO SUMARÍSSIMO/COMUM OU PEÇA CABÍVEL]")
    set_run_font(r, bold=True)
    r = intro.add_run(", pelos fatos e fundamentos a seguir expostos.")
    set_run_font(r)
    add_paragraph(doc)
    add_paragraph(doc)

    add_paragraph(doc, "DA GRATUIDADE DA JUSTICA", "Peticao Titulo Principal")
    add_paragraph(doc, "[Texto objetivo sobre hipossuficiencia e pedido de gratuidade.]")
    add_paragraph(doc)
    add_paragraph(doc, "DOS FATOS", "Peticao Titulo Principal")
    add_paragraph(doc, "[Narrar fatos em ordem cronologica, sempre vinculando documentos relevantes.]")
    add_paragraph(doc)
    add_paragraph(doc, "DOS FUNDAMENTOS JURIDICOS", "Peticao Titulo Principal")
    add_paragraph(doc, "Tese material com fato, fundamento e pedido", "Peticao Subtitulo")
    add_paragraph(doc, "[Abrir o subtópico com fatos e provas relevantes, conectar a norma ou precedente verificado e encerrar com o pedido ou providência concreta decorrente da tese.]")
    add_paragraph(
        doc,
        "[Citacao jurisprudencial ou contratual verificada, com recuo esquerdo de 2,5 cm e fonte 11 pt.]",
        "Peticao Citacao",
    )
    add_paragraph(doc, "[Diante disso, requer a providência específica vinculada a este subtópico.]")
    add_paragraph(doc)
    add_paragraph(doc, "DOS REQUERIMENTOS E PEDIDOS", "Peticao Titulo Principal")
    add_manual_request(doc, "a)", "O ", "recebimento", " da peca com as provas que a instruem.")
    add_manual_request(doc, "b)", "A concessao da ", "gratuidade", " da justica, se cabivel.")
    add_manual_request(doc, "c)", "A ", "procedencia", " do pedido principal, conforme a tese exposta.")
    add_paragraph(doc)
    add_paragraph(doc, "DAS PROVAS", "Peticao Titulo Principal")
    add_paragraph(doc, "Pretende demonstrar a verdade dos fatos por todos os meios de prova em direito admitidos, especialmente prova documental.")
    add_paragraph(doc)
    add_paragraph(doc, "DO VALOR DA CAUSA", "Peticao Titulo Principal")
    add_paragraph(doc, "Atribui-se a causa o valor de R$ [valor].")
    add_paragraph(doc, "[Cidade], [data].")
    add_paragraph(doc)
    add_paragraph(doc, "Carlos Henrique Santana Lima", "Peticao Assinatura Nome")
    add_paragraph(doc, "OAB 60.427/BA", "Peticao Assinatura OAB")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output


def main() -> int:
    import argparse

    parser = argparse.ArgumentParser(description="Generate the peticao-final DOCX template.")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    path = build_template(args.output)
    print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
