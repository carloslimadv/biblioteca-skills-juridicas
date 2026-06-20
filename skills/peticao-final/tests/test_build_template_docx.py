import sys
from pathlib import Path
import pytest
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).parent.parent / "scripts"))
from build_template_docx import set_run_font

def test_set_run_font_default():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("Test")
    set_run_font(run)
    assert run.font.name == "Helvetica"
    assert run.font.size == Pt(12)
    assert run.font.color.rgb == RGBColor(0, 0, 0)
    assert run.bold is None
    assert run.italic is None

    r_pr = run._element.rPr
    assert r_pr is not None
    r_fonts = r_pr.rFonts
    assert r_fonts is not None
    assert r_fonts.get(qn("w:ascii")) == "Helvetica"
    assert r_fonts.get(qn("w:hAnsi")) == "Helvetica"
    assert r_fonts.get(qn("w:cs")) == "Helvetica"

def test_set_run_font_custom():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("Test")
    set_run_font(run, size=14, bold=True, italic=False)
    assert run.font.name == "Helvetica"
    assert run.font.size == Pt(14)
    assert run.bold is True
    assert run.italic is False
